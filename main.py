import streamlit as st
import folium
from streamlit_folium import st_folium
from folium.plugins import Draw
import os
import io
import pyproj
import time
from PIL import Image
from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
import shutil
import arcpy
#from arcpy.da import *
#from arcpy.sa import *
import pathlib
from pathlib import Path
import atexit
import re
import tempfile


class app():
    def __init__(self):
        self.settings()

#Basis for appen
    def settings(self):
        st.set_page_config(page_title='SKREDFAREUTREDNING', layout="wide", page_icon=':snow_capped_mountain:')

#Oppsettet av siden
    def first_page(self):
        st.title('Generer skredfarevurderingsrapport')
        st.markdown('Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.')
#Interaktivt kart som ligger i egen metode under
        self.interactive_map()
        with st.container():
#Input for kartproduksjon og rapporten
            st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
            self.norsk = st.radio('Velg skriftsspråk: ', ['Bokmål','Nynorsk'])
            st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
            self.s = st.radio('Hvilken sikkerhetsklasse skal det utredes for: ', ['S2','S3'])
            #kart = ['Kronedekning', 'Skogtype', 'Markfuktighet', 'Berggrunn', 'Faresoner m dim.skred', 'Granada', 'Helning', 'Skygge', 'Drenering','JordFlom_akt', 'Løsmasser', 'Rapport', 'Sikringstiltak', 'Snoskred_akt', 'Steinsprang_akt', 'Oversiktskart', 'Ortofoto', 'V_Faresoner A3', 'V_Helning A3', 'V_Skog A3', 'V_Registreringskart A3']
            #self.velg_kart = st.multiselect('Velg hvilke kart du ønsker: ', kart, kart[:])
            self.sted = st.text_input('Stedsnavn: ')
            self.tif_file = st.file_uploader('Last opp høydemodell (.tif): ', type=['tif','tiff'])
            self.zoom_events = st.number_input('Zoom for skredhendelseskartet: ', 0, None, 5)
            self.zoom_overview = st.number_input('Zoom for oversiktskartene: ', 0, None, 20)
            self.oppdragsnummer = st.text_input('Oppdragsnummer: ')
            self.dato = st.date_input('Dato: ')
            self.av = st.text_input('Hvem har skrevet rapporten: ')
            self.ks = st.text_input('Hvem har gjort KS av rapporten: ')
            self.oppgiver = st.text_input('Oppdragsgiver: ')
            st.write('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
            self.flomvei_radio = st.toggle('Gjennomfør flomveisanalyse :ocean: (Dette vil ta noen minutter lengre tid)')


            save_path = Path(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS', self.tif_file.name)
            with open(save_path, mode='wb') as w:
                w.write(self.tif_file.getvalue())


#Starter produskjon av kart ved klikk på knapp ved metoden run_maps()
        if st.button('Generer kart'):
            if self.out['all_drawings'] != None:
                self.execute_function_when_filled()

            else:
                st.markdown('Tegn kartleggings-og påvikningsområdet før du klikker "Generer kart"')



#Metoden for det interaktive kartet
    #tiles = 'https://services.geodataonline.no/arcgis/rest/services/Geocache_UTM33_EUREF89/GeocacheBilder/MapServer', attr = "Kilde: geodata.no"
    def interactive_map(self):
        m = folium.Map(location=[65, 11], zoom_start=5, tiles='https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',attr = "Tiles &copy; Esri &mdash; Source: Esri, i-cubed, USDA, USGS, AEX, GeoEye, Getmapping, Aerogrid, IGN, IGP, UPR-EGP, and the GIS User Community")

        Draw(export=False,
             draw_options={'polyline': {'showLength': True, 'metric': True, 'feet': False},
                              'rectangle': False,
                              'polygon': {'showArea': True, 'showLength': False, 'metric': False, 'feet': False},
                              'circle': False,
                              'circlemarker': False,
                              'marker': False,
                              },
                edit_options={'poly': {'allowIntersection': False}}).add_to(m)

        map_style = folium.raster_layers.TileLayer(tiles='https://{s}.tile.thunderforest.com/landscape/{z}/{x}/{y}.png?apikey=a129367b3b1146c794437de8b3d50388',
            attr='&copy; <a href="http://www.thunderforest.com/">Thunderforest</a>, &copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors')
        map_style.add_to(m)

        folium.raster_layers.WmsTileLayer(url=r"https://geo.ngu.no/mapserver/LosmasserWMS2?request=GetCapabilities&service=WMS",
                                          layers= "Losmasse_flate",
                                          name='NGU Løsmassekart',
                                          attr='NGU',
                                          show=False
                                          ).add_to(m)

        folium.raster_layers.WmsTileLayer(url=r"https://geo.ngu.no/mapserver/BerggrunnWMS3?request=GetCapabilities&SERVICE=WMS",
                                          layers= "Berggrunn_nasjonal_hovedbergarter",
                                          name='Nasjonal berggrunn',
                                          attr='NGU',
                                          show=False
                                          ).add_to(m)



        lc = folium.map.LayerControl(position='topright',collapsed=True)
        lc.add_to(m)
        self.out = st_folium(m, width=800, height=700)

        av_logo = Image.open(r'C:\Users\niklas.brede\pythonProject\SKREDMAL\av-logo.png')
        st.sidebar.image(av_logo, caption=None, width=250)
        st.sidebar.header('Objekter for kartproduskjon i ArcGIS (EU89 UTM 33N)')
        st.sidebar.write('Koordinater for kartomriss')

        self.utm_north = self.lat_long_to_utm(self.show_extent(self.out)[2], self.show_extent(self.out)[3])
        self.utm_south = self.lat_long_to_utm(self.show_extent(self.out)[0], self.show_extent(self.out)[1])

        st.sidebar.write('Nord-østlige hjørne:\n ', self.utm_north)
        st.sidebar.write('Sør-vestlige hjørne:\n ', self.utm_south)

        st.sidebar.write(
            'Tegn polygon for kartleggingsområdet og påvirkingsområdet i kartet og klikk på knappen under. Tegn kun to polygon i kartet og tegn kartleggingsområdet FØRST.')

        if st.sidebar.button("Vis koordinater til kartleggings-og påvirkningsområdet"):

            if self.draw_areas(self.out) != False:
                st.sidebar.write('Koordinater for kartleggingsområdet: ', self.draw_areas(self.out)[0])
                st.sidebar.write('Koordinater for påvirkningsområdet: ', self.draw_areas(self.out)[1])

            else:
                st.sidebar.write('Tegn ett karleggingsområdet og ett påvirkningsområdet som polygon.')


#Henter ut koordinatene fra kartomrisset til kartet
    def show_extent(self, output):
        swx = output['bounds']['_southWest']['lat']
        swy = output['bounds']['_southWest']['lng']
        nex = output['bounds']['_northEast']['lat']
        ney = output['bounds']['_northEast']['lng']

        return swx, swy, nex, ney


#Henter ut koordinatene til polygonene tegnet i kartet
    def draw_areas(self, output):

        if len(output['all_drawings']) != 2:
            return False
        else:
            krt = output['all_drawings'][0]['geometry']['coordinates']
            pvik = output['all_drawings'][1]['geometry']['coordinates']

            krt_33 = []
            pvik_33 = []

            for inner_list in krt:
                for sublist in inner_list:
                    x = float(sublist[1])
                    y = float(sublist[0])
                    new = self.lat_long_to_utm(x, y)
                    krt_33.append(new)

            for inner_list in pvik:
                for sublist in inner_list:
                    x = float(sublist[1])
                    y = float(sublist[0])
                    new = self.lat_long_to_utm(x, y)
                    pvik_33.append(new)

            krt_33.pop()
            pvik_33.pop()

            return krt_33, pvik_33


#Konverterer koordinater fra lat/long til WGS84 UTM 33N
    def lat_long_to_utm(self, lat, lon):
        # Define the coordinate systems using EPSG codes for simplicity
        proj_wgs84 = pyproj.Proj(proj='latlong', datum='WGS84')
        proj_utm = pyproj.Proj(proj='utm', zone=33, datum='WGS84')

        transformer = pyproj.Transformer.from_proj(proj_wgs84, proj_utm, always_xy=True)

        utm_easting, utm_northing = transformer.transform(lon, lat)

        return utm_easting, utm_northing

#Sjekker antall filer i mappen hvert 5 sek og kaller metoden 'report()' når antallet er nådd
    def execute_function_when_filled(self):
        with st.spinner('Jobber med saken, dette vil ta flere minutter :coffee:'):
            MakeMaps(bkmk='MyBookmark', tif_file=self.tif_file.name, zoom_events=self.zoom_events,
                     zoom_overview=self.zoom_overview, oppdragsnummer=self.oppdragsnummer, dato=self.dato,
                     av=self.av, ks=self.ks, oppgiver=self.oppgiver,
                     polygons=str(self.draw_areas(self.out)), extent=str(self.utm_south+self.utm_north), flom=self.flomvei_radio)
            self.report()


#Setter sammen rapport med kartene og setter inn tekst fra input
    def report(self):
        dir = r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\APP'


#Skiller mellom 4 maler, en for S2 og en for S3, en for bokmål og en for nynorsk
        def velg_rapport(mal):
            word_file = os.path.join(dir, mal)
            dir_pic = r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\eksporterte_kart'
            document = Document(word_file)
            styles = document.styles
            style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)

            # Setter inn kartene og input på steder i malen som har en spesiell tekst
            def insert_jpg_into_word(word_string, print_map):
                for paragraph_index, paragraph in enumerate(document.paragraphs):
                    if word_string in paragraph.text:
                        map_line = paragraph_index
                        break

                document.paragraphs[map_line].clear()
                run_kart = document.paragraphs[map_line].add_run()
                kart = run_kart.add_picture(os.path.join(dir_pic, print_map))
                kart.width = Cm(16)
                kart.height = Cm(16)

            def insert_txt(eks_str, ny_str):
                for paragraph_index, paragraph in enumerate(document.paragraphs):
                    if eks_str in paragraph.text:
                        linje_til_motstand = paragraph_index
                        innhold = document.paragraphs[linje_til_motstand].text
                        document.paragraphs[linje_til_motstand].clear()
                        document.paragraphs[linje_til_motstand].text = innhold.replace(eks_str, ny_str)

                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if eks_str in cell.text:
                                cell.text = cell.text.replace(eks_str, ny_str)

            insert_jpg_into_word('[BERGGRUNN]','Berggrunn.jpg')
            insert_jpg_into_word('[SKREDHENDELSER]', 'Skredhendelser.jpg')
            insert_jpg_into_word('[KRONEDEKNING]', 'Kronedekning.jpg')
            insert_jpg_into_word('[LØSMASSER]', 'Løsmasser.jpg')
            insert_jpg_into_word('[SKOGTYPE]', 'Skogtype.jpg')
            insert_jpg_into_word('[MARKFUKTGHET]','Markfuktighet.jpg')
            insert_jpg_into_word('[JORDFLOM_AKT]', 'JordFlom_akt.jpg')
            insert_jpg_into_word('[SNOSKRED_AKT]', 'Snoskred_akt.jpg')
            insert_jpg_into_word('[STEINSPRANG_AKT]', 'Steinsprang_akt.jpg')
            insert_jpg_into_word('[OVERSIKTSKART]', 'Oversiktskart.jpg')
            insert_jpg_into_word('[HELNING]', 'Helning.jpg')
            insert_jpg_into_word('[ORTOFOTO]', 'Ortofoto.jpg')
            insert_jpg_into_word('[BERGGRUNN]', 'Berggrunn.jpg')
            if self.flomvei_radio == True:
                insert_jpg_into_word('[DRENERING]', 'Drenering.jpg')
                insert_jpg_into_word('[SKYGGE]', 'Skygge.jpg')
            insert_jpg_into_word('[V_SKOG]', 'V_Skog A3.jpg')
            insert_jpg_into_word('[V_HELNING]', 'V_Helning A3.jpg')

            insert_txt('[OPPDRAGSGIVER]', self.oppgiver)
            insert_txt('[DATO]', str(self.dato))
            insert_txt('[SKREVET_AV]', self.av)
            insert_txt('[KS_GJORT_AV]', self.ks)
            insert_txt('[OPPDRAGSNUMMER]', str(self.oppdragsnummer))
            insert_txt('[STED]', self.sted)

            bio = io.BytesIO()
            document.save(bio)

            def zip_directory_exclude_lock(src_dir, zip_file):
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Walk the source directory
                    for root, dirs, files in os.walk(src_dir):
                        for file in files:
                            if not file.endswith(".lock"):
                                # Compute the source file path
                                src_file = os.path.join(root, file)
                                # Compute the destination file path
                                dest_file = os.path.join(temp_dir, os.path.relpath(src_file, src_dir))
                                # Create the destination directory if it doesn't exist
                                os.makedirs(os.path.dirname(dest_file), exist_ok=True)
                                # Copy the file to the temporary directory
                                shutil.copy2(src_file, dest_file)

                    # Create the zip file from the temporary directory
                    shutil.make_archive(zip_file, 'zip', temp_dir)


# Når dokumentet er lagret blir det zippet sammen med ArcGIS-prosjektet og det dukker det opp en knapp som gjør det mulig å laste ned alt i en zippet fil
            if document:
                with open(
                        f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\RAPPORT\\Skredfarerapport- {self.sted}.docx',
                        'wb') as outputfile:
                    outputfile.write(bio.getvalue())

                if os.path.exists(
                        f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\RAPPORT\\Skredfarerapport- {self.sted}.docx'):

                    zip_directory_exclude_lock(f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS',f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\Skredfarerapport- {self.sted}')

                    if os.path.exists(
                            f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\Skredfarerapport- {self.sted}.zip'):
                        st.balloons()

                        st.write('Prosjektet er ferdig, last ned zippet fil med rapport og ArcGIS Pro-prosjektmappen.')

#Funksjonen sletter midlertidig data og reloader siden etter zip-filen er lastet ned
                        def reset():
                            time.sleep(4)
                            for filename in os.listdir(
                                    r'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\eksporterte_kart'):
                                if os.path.isfile(os.path.join(
                                        r'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\eksporterte_kart',
                                        filename)):
                                    os.remove(os.path.join(r'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\eksporterte_kart', filename))

                            if os.path.isfile(
                                    f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\Skredfarerapport- {self.sted}.zip'):
                                os.remove(f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\Skredfarerapport- {self.sted}.zip')
                                os.remove(f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\GIS\\RAPPORT\\Skredfarerapport- {self.sted}.docx')

                            if os.path.isfile(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Kartleggingsområdet.shp'):
                                os.remove(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Kartleggingsområdet.shp')

                            if os.path.isfile(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Påvirkningsområdet.shp'):
                                os.remove(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Påvirkningsområdet.shp')

                            if os.path.isfile(os.path.join(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS', self.tif_file.name)):
                                os.remove(os.path.join(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS', self.tif_file.name))

                            st.rerun()



                        with open(f'S:\\Felles\\SamferdselInfrastruktur\\skredmal_streamlit\\Skredfarerapport- {self.sted}.zip', 'rb') as fp:

                            st.download_button(
                                label="Last ned ZIP",
                                data=fp,
                                file_name=f"Skredfarerapport- {self.sted}.zip",
                                mime="application/zip", on_click=reset)

        if self.s == 'S2' and self.norsk == 'Nynorsk':
            velg_rapport('RIGberg_Skredfareutgreiing_nynorsk_S2')
        elif self.s == 'S2' and self.norsk == 'Bokmål':
            velg_rapport('RIGberg_Skredfareutredning_bokmål_S2')
        elif self.s == 'S3' and self.norsk == 'Nynorsk':
            velg_rapport('RIGberg_Skredfareutgreiing_nynorsk_S3')
        elif self.s == 'S3' and self.norsk == 'Bokmål':
            velg_rapport('RIGberg_Skredfareutredning_bokmål_S3')


    def main(self):
        self.first_page()




#Denne klassen produserer kartene ved bruk av arcpy
class MakeMaps():

    def __init__(self, bkmk, tif_file, zoom_events, zoom_overview, oppdragsnummer, dato, av, ks, oppgiver, polygons, extent, flom):
        self.bkmk = bkmk
        self.tif_file = tif_file
        self.scale = 0.9
        self.zoom_events = zoom_events
        self.zoom_overview = zoom_overview
        self.oppdragsnummer = oppdragsnummer
        self.dato = dato
        self.av = av
        self.ks = ks
        self.oppgiver = oppgiver
        self.polygons = polygons
        self.extent = extent
        self.flom = flom

        input_list = polygons

        lst = self.parse_list_of_tuples(input_list)
        first_list = []
        second_list = []
        current_list = first_list

        for item in lst:
            if item == ']':
                current_list = second_list
                continue
            current_list.append(item)

        string1 = ''.join(first_list)
        numbers1 = re.findall(r'\d+\.\d+|\d+', string1)
        numbers1 = list(filter(None, numbers1))

        string2 = ''.join(second_list)
        numbers2 = re.findall(r'\d+\.\d+|\d+', string2)
        numbers2 = list(filter(None, numbers2))

        self.make_bookmark(extent=self.extent)
        self.add_polygons(numbers1, numbers2)
        time.sleep(3)

        for maps in aprx.listMaps():
            pa = maps.addDataFromPath(os.path.join(directory, r'Påvirkningsområdet.shp'))
            symbologyLayer_pa = os.path.join(directory, 'LYR//pvirk.lyrx')
            arcpy.management.ApplySymbologyFromLayer(pa, symbologyLayer_pa)
            ka = maps.addDataFromPath(os.path.join(directory, r'Kartleggingsområdet.shp'))
            symbologyLayer_ka = os.path.join(directory, 'LYR//kartleggingsom.lyrx')
            arcpy.management.ApplySymbologyFromLayer(ka, symbologyLayer_ka)

        if flom == False:
            self.all_but_runoff()
        else:
            self.all()
        arcpy.ClearWorkspaceCache_management()

        # Prosessering av strengene med koordinater til rett format
    def parse_list_of_tuples(self,list_str):
        # Convert strings to lists of floats
        result = []
        for i in list_str:
            for j in i:
                if j.isdigit() or '.':
                    result.append(j)
                else:
                    pass

        return result

    # Bruker koordinatene til kartleggings-og påvirkningsområdet for å lage SHP-filer
    def add_polygons(self,numbers1,numbers2):
        krt = []
        for i in range(0, len(numbers1), 2):
            krt.append([float(numbers1[i]), float(numbers1[i + 1])])

        pvirk = []
        for i in range(0, len(numbers2), 2):
            pvirk.append([float(numbers2[i]), float(numbers2[i + 1])])

        krt_shp = r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Kartleggingsområdet.shp'
        pvirk_shp = r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Påvirkningsområdet.shp'
        infc = os.path.join(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS', self.tif_file)

        spatialref = arcpy.Describe(infc).spatialReference

        if arcpy.Exists(krt_shp):
            arcpy.Delete_management(krt_shp)

        if arcpy.Exists(pvirk_shp):
            arcpy.Delete_management(pvirk_shp)

        arcpy.CreateFeatureclass_management(*os.path.split(krt_shp), geometry_type="Polygon")
        arcpy.CreateFeatureclass_management(*os.path.split(pvirk_shp), geometry_type="Polygon")

        icursor_k = arcpy.da.InsertCursor(krt_shp, ["SHAPE@"])
        icursor_p = arcpy.da.InsertCursor(pvirk_shp, ["SHAPE@"])

        array = []
        for feature in krt:
            line = arcpy.Point(feature[0], feature[1])
            array.append(line)

        array = arcpy.Array(array)
        array.append(array[0])
        kartleggingsom = arcpy.Polygon(array, spatialref)

        array2 = []
        for feature in pvirk:
            line2 = arcpy.Point(feature[0], feature[1])
            array2.append(line2)

        array2 = arcpy.Array(array2)
        array2.append(array2[0])
        pavirkningsom = arcpy.Polygon(array2, spatialref)

        icursor_k.insertRow([kartleggingsom])
        icursor_p.insertRow([pavirkningsom])

    # Lager et bookmark fra koordinatene til kartomrissert fra streamlit
    def make_bookmark(self,extent):

        extent = extent.split(',')

        xmin = float(extent[3].replace(')', ''))
        ymin = float(extent[0].replace('(', ''))
        xmax = float(extent[1])
        ymax = float(extent[2])
        ext = arcpy.Extent(ymin, xmin, ymax, xmax)

        bookmark_name = "MyBookmark"

        lyt = aprx.listLayouts()[17]
        mf = lyt.listElements('MAPFRAME_ELEMENT')[0]

        try:
            mf.removeBookmark("MyBookmark")
        except:
            pass

        mf.camera.setExtent(ext)
        mf.createBookmark(bookmark_name)

    def _to_bookmark(self):
        for m in aprx.listMaps():
            for i in m.listBookmarks():
                if i.name == self.bkmk:
                    lyt = aprx.listLayouts()[17]
                    mf = lyt.listElements('MAPFRAME_ELEMENT')[0]
                    mf.zoomToBookmark(i)
                    extent = mf.camera.getExtent()
                    return extent

    def _export_layout_1frame(self, lyt_name, camera_zoom):
        for i in range(len(aprx.listLayouts())):
            if lyt_name == aprx.listLayouts()[i].name:
                lyt = aprx.listLayouts()[i]
                mf = lyt.listElements('MAPFRAME_ELEMENT')[0]
                mf.camera.setExtent(self._to_bookmark())
                mf.camera.scale *= camera_zoom
                lyt.exportToJPEG(os.path.join(output_directory, lyt_name), 300)

    def _export_layout_2frame(self, lyt_name, overview_zoom, strd_zoom, pdf):
        for i in range(len(aprx.listLayouts())):
            if lyt_name == aprx.listLayouts()[i].name:
                lyt = aprx.listLayouts()[i]
                mf = lyt.listElements('MAPFRAME_ELEMENT')[0]
                mf.camera.setExtent(self._to_bookmark())
                mf.camera.scale *= overview_zoom
                mf2 = lyt.listElements('MAPFRAME_ELEMENT')[1]
                mf2.camera.setExtent(self._to_bookmark())
                mf2.camera.scale *= strd_zoom + 0.3
                lyt.exportToJPEG(os.path.join(output_directory, lyt_name), 300)
                if pdf == True:
                    lyt.exportToPDF(os.path.join(output_directory, lyt_name), 300)

    def _set_symbology(self, path_lyrx, _map):
        symbology_layer = path_lyrx
        new_lyr_file = arcpy.mp.LayerFile(symbology_layer)
        new_lyr = new_lyr_file.listLayers()[0]
        old_lyr = _map.listLayers()[0]
        self._unselect_lyr(_map)
        old_lyr_name = old_lyr.name
        new_lyr.updateConnectionProperties(new_lyr.connectionProperties, old_lyr.connectionProperties)
        new_lyr.name = old_lyr_name
        new_lyr_file.save()
        _map.insertLayer(old_lyr, new_lyr_file)
        _map.removeLayer(old_lyr)

    def _unselect_lyr(self, _map):
        _map.listLayers()[0] = False

    def _move_lyr_to_bottom(self, _map):
        _map.moveLayer(_map.listLayers()[len(_map.listLayers()) - 1], _map.listLayers()[0], 'AFTER')

    def _find_map_index(self, map_str):
        for i in aprx.listMaps():
            if i.name == map_str:
                return i

    def _find_layout_index(self, lyt_str):
        for i in aprx.listLayouts():
            if i.name == lyt_str:
                return i

    def _find_layer_index(self, _map, lyr_str):
        for i in _map.listLayers():
            try:
                if i.name == lyr_str:
                    return i
            except:
                pass

    def _del_data(self, data):
        try:
            arcpy.env.overwriteOutput = True
            arcpy.management.Delete(data)
            pass
        except:
            arcpy.AddError('ArcGIS klarer ikke fjerne mappen "helning" selv, lukk programmet og slett mappen manuelt.')

    def _del_from_legend(self, lyt, leg_str):
        leg = lyt.listElements('LEGEND_ELEMENT')[0]
        for itm in leg.items:
           if itm.name  == leg_str:
                leg.removeItem(itm)

    def _add_to_legend(self, lyt, lyr):
        leg = lyt.listElements('LEGEND_ELEMENT')[0]
        leg.addItem(lyr, 'BOTTOM')

    def _change_text(self, lyt):
        for i in aprx.listLayouts():
            if i.name == lyt:
                for txt in i.listElements('TEXT_ELEMENT'):
                    if txt.name == 'Oppdragsnavn':
                        txt.text = self.oppdragsnummer
                    elif txt.name == 'Dato':
                        txt.text = self.dato
                    elif txt.name == 'Utarbeidet av navn':
                        txt.text = self.av
                    elif txt.name == 'Kontrollert av navn':
                        txt.text = self.ks
                    elif txt.name == 'Oppdragsgiver':
                        txt.text = 'Kartet er utarbeidet av Asplan Viak på oppdrag fra ' + self.oppgiver


    def standard_map(self, _map):
        self._export_layout_1frame(_map, self.scale)

    def event_map(self):
        self._export_layout_1frame('Skredhendelser',self.zoom_events)

    def overview_map(self, _map):
        self._export_layout_2frame(_map,self.zoom_overview,self.scale,pdf=False)

    def hillshade_map(self):
        hillshade_raster = os.path.join(directory, f'Relieff')
        self._del_data(hillshade_raster)
        sky = self._find_map_index('Skygge')
        if os.path.exists(hillshade_raster):
            arcpy.AddError('ArcGIS klarer ikke fjerne mappen "relieff" selv, lukk programmet og slett mappen manuelt.')

        arcpy.env.workspace = directory
        arcpy.env.overwriteOutput = True
        inRaster = self.tif_file

        outHillshade = arcpy.sa.Hillshade(inRaster)
        outHillshade.save(hillshade_raster)
        sky.addDataFromPath(hillshade_raster)

        if sky.listLayers()[0].symbology.colorizer.type == 'RasterUniqueValueColorizer':
            self._set_symbology(os.path.join(directory, r'LYR\Hillshade.lyrx'), sky)

        self._del_from_legend(self._find_layout_index('Skygge'),'Relieff')
        self._add_to_legend(self._find_layout_index('Skygge'),self._find_layer_index(sky,'Relieff'))
        self._export_layout_1frame('Skygge',self.scale)

    def slope_map(self):
        slope_raster = os.path.join(directory, f'Helning')
        self._del_data(slope_raster)
        hel = self._find_map_index('Helning')
        if os.path.exists(slope_raster):
            arcpy.AddError('ArcGIS klarer ikke fjerne mappen "helning" selv, lukk programmet og slett mappen manuelt.')

        arcpy.env.workspace = directory
        arcpy.env.overwriteOutput = True

        inRaster = self.tif_file

        terreng = arcpy.sa.SurfaceParameters(inRaster)
        terreng.save(slope_raster)
        hel.addDataFromPath(slope_raster)

        self._set_symbology(os.path.join(directory,r'LYR\Helning.lyrx'), hel)
        self._move_lyr_to_bottom(hel)
        self._del_from_legend(self._find_layout_index('Helning'),'Helning')
        self._add_to_legend(self._find_layout_index('Helning'),self._find_layer_index(hel,'Helning'))
        self._export_layout_1frame('Helning', self.scale)

    def runoff_map(self):
        self._del_data(os.path.join(directory, 'akkumulering'))
        self._del_data(os.path.join(directory, 'flow_direction.crf'))
        dren = self._find_map_index('Drenering')

        arcpy.env.workspace = directory
        arcpy.env.overwriteOutput = True
        input_raster = self.tif_file

        flow_direction_raster = 'flow_direction.crf'
        flow_accumulation_raster = 'Akkumulering'

        arcpy.gp.FlowDirection_sa(input_raster, flow_direction_raster, '', '', 'MFD')
        arcpy.gp.FlowAccumulation_sa(flow_direction_raster, flow_accumulation_raster, '', 'FLOAT')

        dren.addDataFromPath(os.path.join(directory,flow_accumulation_raster))

        self._set_symbology(os.path.join(directory, r'LYR\MFD-flomvei.lyrx'), dren)
        self._del_from_legend(self._find_layout_index('Drenering'),'Akkumulering')
        self._add_to_legend(self._find_layout_index('Drenering'),self._find_layer_index(dren,'Akkumulering'))
        self._export_layout_1frame('Drenering', self.scale)

    def appendix_map(self, _map):
        self._change_text(_map)
        if _map == 'V_Helning A3':
            self._del_from_legend(self._find_layout_index('V_Helning A3'), 'Helning')
            self._add_to_legend(self._find_layout_index('V_Helning A3'), self._find_layer_index(self._find_map_index('Helning'),'Helning'))
        else:
            pass
        self._export_layout_2frame(_map, self.zoom_overview, self.scale, pdf=True)

    def all(self):
        for i in strd_maps:
            self.standard_map(i)
        self.event_map()
        for i in ovr_maps:
            self.overview_map(i)
        self.slope_map()
        for i in ap_maps:
            self.appendix_map(i)
        self.runoff_map()
        self.hillshade_map()

    def all_but_runoff(self):
        for i in strd_maps:
            self.standard_map(i)
        self.event_map()
        for i in ovr_maps:
            self.overview_map(i)
        self.slope_map()
        for i in ap_maps:
            self.appendix_map(i)
        self.hillshade_map()


if __name__ == "__main__":
    def has_run():
        flag_file_path = r"C:\Users\niklas.brede\PycharmProjects\skredmal\first_run.flag"

        if not os.path.exists(flag_file_path):

            with open(flag_file_path, "w") as flag_file:
                flag_file.write("This flag indicates that the script has been run once.")

            os.system(r'streamlit run C:\Users\niklas.brede\PycharmProjects\skredmal\streamgis.py')

        else:
            pass


    def exit_handler():
        os.remove(os.path.join(r'C:\Users\niklas.brede\PycharmProjects\skredmal', 'first_run.flag'))


    has_run()
    atexit.register(exit_handler)



#Setter opp ArcGIS-filen
    aprx = arcpy.mp.ArcGISProject(r'S:\Felles\SamferdselInfrastruktur\skredmal_streamlit\GIS\Mal_skredfarevurdering.aprx')
    project_name = pathlib.Path(aprx.filePath)
    directory = str(project_name.parent)
    output_directory = os.path.join(directory, r'eksporterte_kart')

    strd_maps = ['Kronedekning', 'Skogtype', 'Markfuktighet', 'Berggrunn', 'JordFlom_akt', 'Løsmasser', 'Snoskred_akt', 'Steinsprang_akt']
    ovr_maps = ['Oversiktskart', 'Ortofoto']
    ap_maps = ['V_Faresoner A3', 'V_Helning A3', 'V_Skog A3']

    a = app()
    a.main()